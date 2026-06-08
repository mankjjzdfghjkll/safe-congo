from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, color=None, italic=False):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Times New Roman"
    run.font.bold  = True
    run.font.size  = Pt(13) if level == 1 else Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def shade_cell(cell, fill_hex="1F497D"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",  "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",   "6"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color","4472C4"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()  # spacing
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSITÉ NOUVEAUX HORIZONS")
set_font(run, size=13, bold=True, color=(0x1F, 0x49, 0x7D))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Faculté des Sciences Informatiques – Génie Logiciel")
set_font(run, size=12, italic=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MAL1471 – Projet Final 2026")
set_font(run, size=14, bold=True, color=(0x1F, 0x49, 0x7D))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Phase 1(a) – Environnement du Stage")
set_font(run, size=13, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rapport soumis par :")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANKAND-A-MUTEB JOSEE")
set_font(run, size=14, bold=True, color=(0x1F, 0x49, 0x7D))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Matricule : SI/20223393")
set_font(run, size=11, italic=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Organisme d'accueil :")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Institut National de Préparation Professionnelle (INPP)")
set_font(run, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Direction Provinciale du Haut-Katanga – Lubumbashi")
set_font(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date de soumission : 27 mai 2026")
set_font(run, size=11, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. L'ORGANISATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1.  L'Organisation")
add_body(doc,
    "L'Institut National de Préparation Professionnelle (INPP) est un établissement "
    "public congolais placé sous la tutelle du Ministère de l'Emploi, du Travail et de "
    "la Prévoyance Sociale. Sa mission principale consiste à assurer la formation "
    "continue, le perfectionnement et la reconversion professionnelle des travailleurs "
    "salariés du secteur formel. L'INPP opère dans le secteur de la formation professionnelle "
    "et s'inscrit dans la politique nationale de développement du capital humain en "
    "République Démocratique du Congo."
)
add_body(doc,
    "En termes de taille et de portée, l'institution dispose d'une présence nationale : "
    "une Direction Générale basée à Kinshasa, des Directions Provinciales dans les "
    "principales provinces du pays, et des antennes locales dans plusieurs villes. "
    "La Direction Provinciale du Haut-Katanga, où j'effectue mon stage, est implantée "
    "à Lubumbashi et couvre l'ensemble de la province minière. Le financement de "
    "l'institution est principalement assuré par des cotisations patronales obligatoires, "
    "comprises entre 0,5 % et 3 % de la masse salariale brute des entreprises, ce qui "
    "lui confère une relative autonomie financière. Cette structure sectorielle m'a "
    "frappée dès mon arrivée : contrairement à une entreprise classique, l'INPP fonctionne "
    "comme un opérateur public de service, répondant à des demandes formulées par des "
    "entreprises partenaires plutôt qu'à un marché commercial."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. MISSION ET STRATÉGIE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.  Mission et Stratégie")
add_body(doc,
    "La mission stratégique que j'ai observée au quotidien se décline en trois axes "
    "concrets : (1) l'adaptation des compétences aux besoins réels des entreprises locales, "
    "notamment dans les secteurs minier et sous-traitant qui dominent l'économie katangaise ; "
    "(2) la modernisation des outils de gestion et de formation ; (3) la digitalisation "
    "progressive des processus internes."
)
add_body(doc,
    "En pratique, lors des réunions d'équipe auxquelles j'ai assisté, les priorités "
    "exprimées par la hiérarchie portaient systématiquement sur la qualité des formations "
    "dispensées et sur l'innovation technique. Un partenariat actif avec la JICA (Agence "
    "Japonaise de Coopération Internationale) est régulièrement cité comme levier "
    "stratégique pour importer de bonnes pratiques en gestion de la formation. Ce contexte "
    "m'a conduite à comprendre que l'INPP ambitionne de devenir un centre de référence "
    "régional en matière de formation technique, ce qui justifie l'importance accordée "
    "à l'informatisation des registres de stagiaires, des évaluations et des certifications."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. STRUCTURE ET ORGANIGRAMME
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3.  Structure et Organigramme")
add_body(doc,
    "La Direction Provinciale du Haut-Katanga s'articule autour de plusieurs services "
    "fonctionnels. Le tableau ci-dessous résume les principaux départements :"
)

# Services table
tbl = doc.add_table(rows=7, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Service", "Rôle principal"]
row0 = tbl.rows[0]
for i, h in enumerate(headers):
    shade_cell(row0.cells[i], "1F497D")
    run = row0.cells[i].paragraphs[0].add_run(h)
    run.font.name  = "Times New Roman"
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data = [
    ("Direction Provinciale",    "Pilotage stratégique, représentation institutionnelle"),
    ("Recouvrement",             "Contrôle et collecte des cotisations patronales"),
    ("Formation / Technique",    "Conception et délivrance des programmes de formation"),
    ("Ressources Humaines",      "Gestion du personnel, recrutement, évaluation"),
    ("Logistique",               "Gestion du matériel, des locaux et des équipements"),
    ("Informatique (Pool IT)",   "Maintenance des systèmes, numérisation, développement applicatif"),
]
for r_idx, (col1, col2) in enumerate(data):
    row = tbl.rows[r_idx + 1]
    for cell in row.cells:
        if (r_idx + 1) % 2 == 0:
            shade_cell(cell, "DCE6F1")
    row.cells[0].paragraphs[0].add_run(col1).font.name = "Times New Roman"
    row.cells[0].paragraphs[0].runs[-1].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(col2).font.name = "Times New Roman"
    row.cells[1].paragraphs[0].runs[-1].font.size = Pt(11)

doc.add_paragraph()

# ─── Org Chart (text-based table) ───────────────────────────────────────────
add_body(doc, "Organigramme simplifié – positionnement de l'équipe Informatique :")

org_tbl = doc.add_table(rows=7, cols=1)
org_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
org_tbl.style = "Table Grid"

org_data = [
    ("Direction Générale – Kinshasa",    "2C5282", True),
    ("▼",                                "FFFFFF", False),
    ("Direction Provinciale – Haut-Katanga",  "2B6CB0", True),
    ("▼",                                "FFFFFF", False),
    ("Service Informatique / Pool IT ◄ MON ÉQUIPE", "2A69AC", True),
    ("▼",                                "FFFFFF", False),
    ("Cellule de Développement & Maintenance ◄ MON POSTE", "3182CE", True),
]

for i, (txt, fill, bold_flag) in enumerate(org_data):
    cell = org_tbl.rows[i].cells[0]
    shade_cell(cell, fill)
    p_cell = cell.paragraphs[0]
    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cell.add_run(txt)
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(11)
    run.font.bold  = bold_flag
    if fill != "FFFFFF":
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        run.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

doc.add_paragraph()
add_body(doc,
    "Mon équipe (Pool Informatique / Cellule de Développement) est rattachée directement "
    "au Service Informatique, lui-même sous l'autorité du Chef de Service, qui rend compte "
    "au Directeur Provincial. Cette ligne hiérarchique courte favorise des échanges "
    "directs et réactifs."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. MON STAGE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4.  Mon Stage et Mon Équipe")
add_body(doc,
    "J'effectue mon stage au sein du Pool Informatique de l'INPP Haut-Katanga, "
    "une cellule fonctionnelle d'environ six personnes (développeurs, techniciens réseau "
    "et administrateurs systèmes). Cette équipe est le centre névralgique de la "
    "transformation numérique de la Direction Provinciale. Sa mission couvre la "
    "maintenance du parc informatique, le développement d'applications métiers internes, "
    "la gestion des bases de données institutionnelles et le support technique aux autres "
    "services."
)
add_body(doc,
    "En termes d'interactions, le Pool IT collabore étroitement avec le service "
    "Formation/Technique pour dématérialiser les dossiers de stagiaires, avec le service "
    "Recouvrement pour automatiser les relances de cotisations, et avec les Ressources "
    "Humaines pour la gestion numérique du personnel. Cette transversalité m'a permis "
    "de comprendre rapidement les flux de données de l'organisation, un avantage "
    "considérable pour la suite du projet MAL1471."
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. MON RÔLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5.  Mon Rôle et Mes Responsabilités")
add_body(doc,
    "Mon titre de poste est Stagiaire Développeuse au sein de la Cellule de "
    "Développement. Mes responsabilités et tâches régulières sont les suivantes :"
)

tasks = [
    ("Développement d'applications internes",
     "Participation à la conception et au codage d'un module de gestion des dossiers "
     "de stagiaires (enregistrement, suivi des présences, génération d'attestations). "
     "J'utilise principalement Python (Flask) pour le back-end et HTML/CSS pour les "
     "interfaces."),
    ("Maintenance et administration de bases de données",
     "Mise à jour régulière de la base de données MySQL centralisant les informations "
     "des stagiaires inscrits, les résultats d'évaluation et les fiches d'entreprises "
     "partenaires."),
    ("Support technique et formation",
     "Assistance aux agents des autres services pour l'utilisation des outils "
     "informatiques existants ; rédaction de mini-guides d'utilisation."),
    ("Analyse des besoins",
     "Participation aux réunions avec les chefs de service pour recueillir leurs "
     "besoins applicatifs et les traduire en spécifications fonctionnelles."),
]

for titre, desc in tasks:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(titre + " : ")
    run_t.font.name = "Times New Roman"
    run_t.font.size = Pt(12)
    run_t.font.bold = True
    run_d = p.add_run(desc)
    run_d.font.name = "Times New Roman"
    run_d.font.size = Pt(12)

doc.add_paragraph()
add_body(doc,
    "Ce rôle m'a confrontée à une réalité opérationnelle importante : la grande partie "
    "des données de l'INPP est encore saisie manuellement sur papier ou dans des "
    "fichiers Excel non standardisés. Cette observation nourrit directement ma réflexion "
    "pour les phases 2 et 3 du projet."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PERSONNES CLÉS – Tableau des parties prenantes
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6.  Personnes Clés – Tableau des Parties Prenantes")

stake_tbl = doc.add_table(rows=6, cols=3)
stake_tbl.style = "Table Grid"
stake_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

s_headers = ["Prénom / Titre", "Poste / Fonction", "Ma relation avec eux"]
s_row0 = stake_tbl.rows[0]
for i, h in enumerate(s_headers):
    shade_cell(s_row0.cells[i], "1F497D")
    run = s_row0.cells[i].paragraphs[0].add_run(h)
    run.font.name  = "Times New Roman"
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

stake_data = [
    ("M. le Directeur Provincial",
     "Directeur Provincial de l'INPP Haut-Katanga",
     "Responsable hiérarchique ultime ; participe aux réunions de cadrage où mes travaux sont présentés."),
    ("M. Kalombo (Chef SI)",
     "Chef du Service Informatique / Encadreur direct de stage",
     "Superviseur immédiat : valide mes livrables, oriente mes tâches et évalue mon travail."),
    ("M. Ilunga (Dev Senior)",
     "Développeur senior – Cellule de Développement",
     "Collègue direct, mentor technique : revoit mon code et m'accompagne sur les modules complexes."),
    ("Mme Numbi (RH)",
     "Responsable des Ressources Humaines",
     "Partie prenante fonctionnelle : exprime les besoins en numérisation des dossiers du personnel."),
    ("M. Kyungu (Formation)",
     "Chef du Service Formation / Technique",
     "Utilisateur final principal des applications que je développe ; source de données de formation."),
]

for r_idx, (col1, col2, col3) in enumerate(stake_data):
    row = stake_tbl.rows[r_idx + 1]
    if (r_idx + 1) % 2 == 0:
        for cell in row.cells:
            shade_cell(cell, "DCE6F1")
    for cell, txt in zip(row.cells, [col1, col2, col3]):
        run = cell.paragraphs[0].add_run(txt)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# RÉFLEXION PERSONNELLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Réflexion Personnelle")
add_body(doc,
    "Ce stage constitue pour moi une immersion très formatrice dans le fonctionnement "
    "d'un établissement public congolais en pleine transition numérique. Plusieurs "
    "constats m'ont marquée : d'abord, l'écart important entre les ambitions de "
    "digitalisation affichées par la direction et la réalité des pratiques quotidiennes "
    "(saisies manuelles, fichiers dispersés, absence d'interconnexion entre les services). "
    "Ensuite, la richesse des données potentiellement exploitables : présences des "
    "stagiaires, résultats d'évaluations, historiques de cotisations, profils d'entreprises. "
    "Ces données existent mais ne sont pas structurées de manière à permettre une analyse "
    "fiable."
)
add_body(doc,
    "Ce que je ne comprends pas encore pleinement, c'est la logique de gouvernance "
    "des données : qui est propriétaire de quelle donnée, et quels services ont "
    "l'autorisation d'y accéder ? Ces questions de droits d'accès et de qualité des "
    "données constitueront un enjeu central lorsque j'aborderai la phase 2 (exploration "
    "des données) et la phase 3 (définition du problème d'apprentissage automatique)."
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
section = doc.sections[0]
footer  = section.footer
fp      = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("MANKAND-A-MUTEB JOSEE  |  MAL1471 – Phase 1(a)  |  INPP Haut-Katanga  |  27 mai 2026")
run.font.name   = "Times New Roman"
run.font.size   = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\PC\Desktop\TP machine learning\MANKAND-A-MUTEB_JOSEE_MAL1471_Phase1a.docx"
doc.save(output_path)
print(f"✅  Document créé avec succès : {output_path}")
